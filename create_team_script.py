from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

print("Creating Team Presentation Script...")

# Create document
doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title
title = doc.add_heading('OKLO: FROM REACTORS TO REVENUE', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(72, 201, 176)

# Subtitle
subtitle = doc.add_paragraph('10-Minute Team Presentation Script')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.bold = True

# Metadata
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Duration: 10 minutes (50 seconds per slide average)\n').font.size = Pt(11)
meta.add_run('Slides: 12\n').font.size = Pt(11)
meta.add_run('Format: Business presentation with team delivery').font.size = Pt(11)

doc.add_paragraph('_' * 100)

# SLIDE 1
doc.add_heading('SLIDE 1: TITLE - "From Reactors to Revenue" [0:00-0:30] (30 sec)', 1)
doc.add_paragraph('[Team Member 1 - Introduction]').runs[0].font.italic = True

script = doc.add_paragraph()
script.add_run('"Good morning/afternoon everyone. Today we\'re presenting ')
script.add_run('"From Reactors to Revenue: Oklo\'s Path to Profitability."').bold = True
script.add_run('\n\nOklo is a nuclear energy company with groundbreaking technology—but like many startups in this space, they\'re facing a critical challenge: ')
script.add_run('How do you generate revenue while waiting years for regulatory approval?').bold = True
script.add_run('\n\nOur team has identified a strategic pivot that could transform Oklo from a pre-revenue reactor builder into a ')
script.add_run('$140 million consulting powerhouse.')
script.add_run('\n\nLet\'s dive in."')

doc.add_paragraph('[ADVANCE TO SLIDE 2]').runs[0].font.italic = True
doc.add_paragraph()

# SLIDE 2
doc.add_heading('SLIDE 2: THE OPPORTUNITY - "80+ Startups Need Help" [0:30-1:15] (45 sec)', 1)
doc.add_paragraph('[Team Member 2 - Problem Setup]').runs[0].font.italic = True

script2 = doc.add_paragraph()
script2.add_run('"The nuclear energy industry is experiencing a renaissance. Over ')
script2.add_run('80 new nuclear startups').bold = True
script2.add_run(' have entered the market in the past five years—companies developing advanced microreactors, small modular reactors, and next-generation designs.\n\nBut here\'s the challenge: ')
script2.add_run('Every single one of them needs to navigate the same regulatory maze that Oklo has already been through.').bold = True
script2.add_run('\n\nThey need licensing expertise. They need safety analysis. They need fuel sourcing strategies. And right now, they\'re scrambling to figure it out on their own."')

doc.add_paragraph('[ADVANCE TO SLIDE 3]').runs[0].font.italic = True
doc.add_paragraph()

# SLIDE 3
doc.add_heading('SLIDE 3: THE CROWDED MARKET [1:15-2:00] (45 sec)', 1)
doc.add_paragraph('[Team Member 1 - Market Context]').runs[0].font.italic = True

script3 = doc.add_paragraph()
script3.add_run('"Look at this market. [POINT TO CHART]\n\nTerraPower has raised $1.4 billion. X-energy, $1 billion. Kairos and Newcleo have raised over $600 million each. That\'s over ')
script3.add_run('$4 billion in capital').bold = True
script3.add_run(' sitting in bank accounts, waiting for NRC approval.\n\nBut funding doesn\'t solve the core problem: ')
script3.add_run('these companies don\'t know how to navigate the regulatory process.').bold = True
script3.add_run('\n\nThey have brilliant engineers and breakthrough technology. What they don\'t have is regulatory experience. That\'s where Oklo comes in."')

doc.add_paragraph('[ADVANCE TO SLIDE 4]').runs[0].font.italic = True
doc.add_paragraph()

# SLIDE 4
doc.add_heading('SLIDE 4: WHAT THEY ALL NEED [2:00-2:45] (45 sec)', 1)
doc.add_paragraph('[Team Member 3 - Paint the Pain]').runs[0].font.italic = True

script4 = doc.add_paragraph()
script4.add_run('"Let me show you what these startups are facing. [POINT TO EACH STAT]\n\n')
script4.add_run('10+ years').bold = True
script4.add_run(' of licensing timeline. That\'s a decade of burning capital with zero revenue.\n\n')
script4.add_run('$50 million-plus').bold = True
script4.add_run(' in regulatory costs—just to get approval. Not to build. Not to operate. Just to get permission.\n\n')
script4.add_run('And here\'s the kicker: Zero reactor experience.').bold = True
script4.add_run(' Most of these startups have never submitted an NRC application. They\'re learning as they go, making expensive mistakes.\n\n[Pause for effect]\n\nThis is ')
script4.add_run('exactly').bold = True
script4.add_run(' the problem Oklo can solve."')

doc.add_paragraph('[ADVANCE TO SLIDE 5]').runs[0].font.italic = True
doc.add_paragraph()

# SLIDE 5
doc.add_heading('SLIDE 5: OKLO HAS DONE THIS [2:45-3:30] (45 sec)', 1)
doc.add_paragraph('[Team Member 2 - The Insight]').runs[0].font.italic = True

script5 = doc.add_paragraph()
script5.add_run('"In 2020, Oklo submitted the first-ever combined license application for an advanced non-light-water reactor. They were pioneers.\n\nIn 2022, the NRC rejected it.\n\nNow, most companies would see that as a failure. We see it as ')
script5.add_run('the most valuable asset Oklo has.').bold = True
script5.add_run('\n\nBecause Oklo now knows:\n• Exactly what the NRC requires\n• What mistakes to avoid\n• How to structure a winning application\n• What documentation is critical\n\nThat knowledge? It\'s worth ')
script5.add_run('$140 million."').bold = True

doc.add_paragraph('[ADVANCE TO SLIDE 6]').runs[0].font.italic = True
doc.add_paragraph()

# SLIDE 6
doc.add_heading('SLIDE 6: THREE REVENUE STREAMS [3:30-4:30] (1 min)', 1)
doc.add_paragraph('[Team Member 1 - The Solution]').runs[0].font.italic = True

script6 = doc.add_paragraph()
script6.add_run('"Here\'s our proposal: ')
script6.add_run('Oklo Technical Services.').bold = True
script6.add_run(' Three distinct revenue streams.\n\n[POINT TO EACH PILLAR]\n\n')
script6.add_run('First: Licensing and Regulatory Consulting.').bold = True
script6.add_run(' Oklo helps startups navigate NRC applications, prepare safety cases, and develop regulatory strategies. Target: ')
script6.add_run('$75 million annually by 2030.').bold = True
script6.add_run('\n\n')
script6.add_run('Second: Safety and Simulation Services.').bold = True
script6.add_run(' Reactor modeling, thermal-hydraulic analysis, probabilistic risk assessments. These require specialized tools and PhD-level expertise. Target: ')
script6.add_run('$42 million annually.').bold = True
script6.add_run('\n\n')
script6.add_run('Third: Fuel Cycle Management.').bold = True
script6.add_run(' HALEU fuel sourcing, recycling strategy, fuel licensing support. This is Oklo\'s unique competitive advantage. Target: ')
script6.add_run('$23 million annually.').bold = True
script6.add_run('\n\nTotal by 2030: ')
script6.add_run('$140 million in annual consulting revenue."').bold = True

doc.add_paragraph('[ADVANCE TO SLIDE 7]').runs[0].font.italic = True
doc.add_paragraph()

# SLIDE 7
doc.add_heading('SLIDE 7: MARKET OPPORTUNITY [4:30-5:15] (45 sec)', 1)
doc.add_paragraph('[Team Member 3 - Market Sizing]').runs[0].font.italic = True

script7 = doc.add_paragraph()
script7.add_run('"Let\'s talk about market size. [POINT TO FUNNEL]\n\n')
script7.add_run('$6.2 billion').bold = True
script7.add_run(' global total addressable market. That\'s 80+ startups times their average lifetime licensing spend.\n\n')
script7.add_run('$1.9 billion').bold = True
script7.add_run(' U.S. serviceable market. We\'re focusing on U.S. and allied nations where Oklo can realistically operate.\n\n')
script7.add_run('$570 million').bold = True
script7.add_run(' is our target. That represents capturing 30% market share of the advanced reactor consulting segment by 2030.\n\nThis is a ')
script7.add_run('massive, growing market with limited competition."').bold = True

doc.add_paragraph('[ADVANCE TO SLIDE 8]').runs[0].font.italic = True
doc.add_paragraph()

# SLIDE 8
doc.add_heading('SLIDE 8: 5-YEAR GROWTH [5:15-6:00] (45 sec)', 1)
doc.add_paragraph('[Team Member 2 - Growth Trajectory]').runs[0].font.italic = True

script8 = doc.add_paragraph()
script8.add_run('"Now let\'s look at the growth trajectory. [POINT TO CHART]\n\nThis chart shows our projected revenue growth over five years.\n\nYear 1 (2026): $4.5 million. Modest start, proof of concept.\n\nYear 3 (2028): $35 million with expanding client base.\n\nYear 5 (2030): ')
script8.add_run('$140 million in annual revenue.').bold = True
script8.add_run('\n\nNotice the curve accelerates. That\'s because:\n• Licensing services scale quickly (high demand)\n• Safety services build on repeatable models\n• Fuel services create long-term contracts\n\nThis is ')
script8.add_run('real, recurring, high-margin revenue."').bold = True

doc.add_paragraph('[ADVANCE TO SLIDE 9]').runs[0].font.italic = True
doc.add_paragraph()

# SLIDE 9
doc.add_heading('SLIDE 9: PROFITABILITY [6:00-6:45] (45 sec)', 1)
doc.add_paragraph('[Team Member 1 - Financial Model]').runs[0].font.italic = True

script9 = doc.add_paragraph()
script9.add_run('"Let\'s talk profitability.\n\n')
script9.add_run('$42 million in EBITDA by 2030.').bold = True
script9.add_run(' That\'s 35-40% margin—exceptional for a professional services business.\n\nWhy such high margins?\n• Leverages existing Oklo expertise (low marginal cost)\n• Premium pricing (specialized knowledge commands high rates)\n• Scalable delivery model (repeatable processes)\n\nConsulting firms in this space charge ')
script9.add_run('$250-500 per hour').bold = True
script9.add_run(' for nuclear regulatory expertise. Oklo can command the top end of that range because they\'ve actually ')
script9.add_run('been through the process."').bold = True

doc.add_paragraph('[ADVANCE TO SLIDE 10]').runs[0].font.italic = True
doc.add_paragraph()

# SLIDE 10
doc.add_heading('SLIDE 10: 4-YEAR ROADMAP [6:45-7:30] (45 sec)', 1)
doc.add_paragraph('[Team Member 3 - Implementation]').runs[0].font.italic = True

script10 = doc.add_paragraph()
script10.add_run('"Here\'s how we get there. [POINT TO TIMELINE]\n\n')
script10.add_run('2026 - Launch Phase:').bold = True
script10.add_run(' Hire core team, sign first 3 pilot clients, establish service offerings.\n\n')
script10.add_run('2027 - Scale Phase:').bold = True
script10.add_run(' Expand to 20 clients, build case studies, refine delivery model.\n\n')
script10.add_run('2028 - Expansion Phase:').bold = True
script10.add_run(' Reach 40 clients, establish market leadership, launch fuel cycle services.\n\n')
script10.add_run('2030 - Market Leader:').bold = True
script10.add_run(' 100+ clients, $140M annual revenue, dominant position in advanced reactor consulting.\n\nThis is ')
script10.add_run('aggressive but achievable').bold = True
script10.add_run('—we\'re leveraging existing relationships and Oklo\'s brand credibility."')

doc.add_paragraph('[ADVANCE TO SLIDE 11]').runs[0].font.italic = True
doc.add_paragraph()

# SLIDE 11
doc.add_heading('SLIDE 11: THE INVESTMENT [7:30-8:30] (1 min)', 1)
doc.add_paragraph('[Team Member 2 - The Ask]').runs[0].font.italic = True

script11 = doc.add_paragraph()
script11.add_run('"So what does this require?\n\n')
script11.add_run('Initial investment: $2 million.').bold = True
script11.add_run('\n\nThat covers:\n• Hiring 3-5 senior consultants (licensing experts, safety engineers)\n• Software tools and modeling platforms\n• Marketing and business development\n• Legal entity setup and compliance\n\n')
script11.add_run('Breakeven: Year 2 (2027).').bold = True
script11.add_run(' That\'s remarkably fast.\n\n')
script11.add_run('Return on investment: 70x by Year 5.').bold = True
script11.add_run('\n\nBut beyond the numbers, this delivers ')
script11.add_run('strategic benefits:').bold = True
script11.add_run('\n• Revenue diversification (reduces dependence on Aurora deployment)\n• Market intelligence (see which technologies are winning)\n• Ecosystem leadership (position Oklo as the industry enabler)\n• Talent utilization (make experts billable during Aurora construction)\n• Future fuel customers (today\'s consulting clients become tomorrow\'s fuel buyers)"')

doc.add_paragraph('[ADVANCE TO SLIDE 12]').runs[0].font.italic = True
doc.add_paragraph()

# SLIDE 12
doc.add_heading('SLIDE 12: CONCLUSION [8:30-10:00] (1 min 30 sec)', 1)
doc.add_paragraph('[Team Member 1 - Closing]').runs[0].font.italic = True

script12 = doc.add_paragraph()
script12.add_run('"Let me bring this home.\n\nOklo is sitting on ')
script12.add_run('the most valuable failure in the nuclear industry.').bold = True
script12.add_run(' They know what 80 other startups desperately need to learn.\n\nRight now, those startups are burning through capital, making mistakes, and facing years of regulatory delays. Oklo can ')
script12.add_run('help them succeed faster').bold = True
script12.add_run(' while ')
script12.add_run('generating immediate revenue.').bold = True
script12.add_run('\n\n[PAUSE]\n\nThe choice is clear:\n\nWait 7+ years for Aurora to generate revenue...\n\nOr start making ')
script12.add_run('$140 million annually').bold = True
script12.add_run(' in ')
script12.add_run('six months').bold = True
script12.add_run(' while building Aurora in parallel.\n\n[PAUSE, make eye contact]\n\nThis isn\'t just a consulting business. It\'s ')
script12.add_run('Oklo becoming the backbone of the nuclear renaissance.').bold = True
script12.add_run('\n\nThank you. We\'re ready for questions."')

doc.add_paragraph()
doc.add_paragraph('[Team: Smile, open for Q&A]').runs[0].font.italic = True

doc.add_page_break()

# Q&A SECTION
doc.add_heading('Q&A PREPARATION', 1)

qa_list = [
    {
        'q': '"Won\'t consulting distract from building Aurora?"',
        'a': '"Actually, it complements it. Oklo\'s licensing experts are currently overhead during Aurora\'s construction phase. By making them billable, we generate revenue AND keep top talent engaged. Plus, customer intelligence from consulting directly informs Aurora\'s strategy. It\'s synergistic, not distracting."',
        'who': 'Team Member 2'
    },
    {
        'q': '"Why would competitors hire Oklo?"',
        'a': '"They\'re not really competitors—most serve different market segments. A mining microreactor is very different from Oklo\'s data center focus. Plus, they\'ll hire consultants anyway. Better that Oklo captures that revenue and builds relationships. Today\'s consulting client could be tomorrow\'s fuel customer."',
        'who': 'Team Member 3'
    },
    {
        'q': '"Doesn\'t Oklo\'s 2022 rejection hurt credibility?"',
        'a': '"It\'s actually our biggest strength. Oklo learned exactly what the NRC wants. Startups want consultants who\'ve been through the process—including failure. That\'s more valuable than theoretical expertise. Our pitch is: \'We learned the hard way so you don\'t have to.\'"',
        'who': 'Team Member 1'
    },
    {
        'q': '"Is the market really big enough?"',
        'a': '"80 startups, each spending $5-20M on licensing over 5-7 years equals a $1.9B serviceable market. We\'re targeting 30% share. Even if half the startups fail, it\'s still a massive opportunity. And we can expand internationally—Canada, UK, and EU are all developing advanced reactor programs."',
        'who': 'Team Member 2'
    },
    {
        'q': '"How quickly can you launch?"',
        'a': '"Six months to first revenue. Oklo already has the experts on staff. We need to hire 3-5 additional people, define service packages, and launch a pilot with 3 friendly clients. That\'s Months 1-6. Revenue starts flowing immediately after."',
        'who': 'Team Member 3'
    }
]

for qa in qa_list:
    q_para = doc.add_paragraph()
    q_para.add_run('Q: ').bold = True
    q_para.add_run(qa['q']).italic = True

    a_para = doc.add_paragraph()
    a_para.add_run(f"A ({qa['who']}): ").bold = True
    a_para.add_run(qa['a'])

    doc.add_paragraph()

doc.add_page_break()

# DELIVERY TIPS
doc.add_heading('TEAM DELIVERY TIPS', 1)

doc.add_heading('Speaker Transitions:', 2)
doc.add_paragraph('• Rehearse handoffs between team members')
doc.add_paragraph('• Make eye contact when passing the presentation')
doc.add_paragraph('• Use names: "Now I\'ll turn it over to [Name] to discuss..."')

doc.add_heading('Timing:', 2)
doc.add_paragraph('• Practice with a timer—hit 10 minutes exactly 3 times before presenting')
doc.add_paragraph('• Each person should know their section timing')
doc.add_paragraph('• Have a timekeeper give subtle signals (2 min warning, 1 min warning)')

doc.add_heading('Visual Engagement:', 2)
doc.add_paragraph('• Point to charts and graphics when referencing them')
doc.add_paragraph('• Use hand gestures for numbers (count on fingers for "three revenue streams")')
doc.add_paragraph('• Make eye contact with different audience members')

doc.add_heading('Energy & Pacing:', 2)
doc.add_paragraph('• Start strong (hook them in first 30 seconds)')
doc.add_paragraph('• Vary vocal tone—don\'t monotone')
doc.add_paragraph('• Use pauses for emphasis (after key numbers)')
doc.add_paragraph('• End powerfully (make the close memorable)')

doc.add_heading('Team Coordination:', 2)
doc.add_paragraph('• Stand/sit in designated spots')
doc.add_paragraph('• Non-speaking members should look attentive (no fidgeting)')
doc.add_paragraph('• Support each other (nod when teammate makes a good point)')
doc.add_paragraph('• For Q&A: whoever knows the answer best should respond')

# TIMING TABLE
doc.add_page_break()
doc.add_heading('TIMING BREAKDOWN', 1)

table = doc.add_table(rows=13, cols=4)
table.style = 'Light Grid Accent 1'

# Header
hdr = table.rows[0].cells
hdr[0].text = 'Slide'
hdr[1].text = 'Speaker'
hdr[2].text = 'Duration'
hdr[3].text = 'Cumulative'

# Data
data = [
    ('1 - Title', 'Member 1', '0:30', '0:30'),
    ('2 - Opportunity', 'Member 2', '0:45', '1:15'),
    ('3 - Market', 'Member 1', '0:45', '2:00'),
    ('4 - Needs', 'Member 3', '0:45', '2:45'),
    ('5 - Oklo\'s Edge', 'Member 2', '0:45', '3:30'),
    ('6 - Three Streams', 'Member 1', '1:00', '4:30'),
    ('7 - Market Size', 'Member 3', '0:45', '5:15'),
    ('8 - Growth', 'Member 2', '0:45', '6:00'),
    ('9 - Profitability', 'Member 1', '0:45', '6:45'),
    ('10 - Roadmap', 'Member 3', '0:45', '7:30'),
    ('11 - Investment', 'Member 2', '1:00', '8:30'),
    ('12 - Conclusion', 'Member 1', '1:30', '10:00')
]

for idx, (slide, speaker, duration, cumulative) in enumerate(data, 1):
    row = table.rows[idx]
    row.cells[0].text = slide
    row.cells[1].text = speaker
    row.cells[2].text = duration
    row.cells[3].text = cumulative

doc.add_paragraph()

# Final message
doc.add_paragraph()
closing = doc.add_paragraph('READY TO PRESENT!')
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
closing.runs[0].bold = True
closing.runs[0].font.size = Pt(16)
closing.runs[0].font.color.rgb = RGBColor(72, 201, 176)

message = doc.add_paragraph(
    'Your team has a powerful, data-driven story. '
    'Deliver it with confidence and coordination.\n'
    'Practice together, support each other, and nail this presentation!'
)
message.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save
output_file = 'OKLO_Team_Presentation_Script_10min.docx'
doc.save(output_file)

print(f"SUCCESS! Team script created: {output_file}")
print("  12 slides")
print("  Exactly 10 minutes")
print("  3 team members")
print("  Includes Q&A prep and delivery tips")
